#!/usr/bin/env python3
"""Build a styled .docx meeting minutes document from a structured JSON file.

Usage:
    python3 build_docx.py <input.json> <output.docx>

See SKILL.md for the JSON schema this expects.
"""
import json
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

NAVY = RGBColor(0x1F, 0x38, 0x64)
GRAY = RGBColor(0x80, 0x80, 0x80)
BORDER_GRAY = "BFBFBF"
LABEL_SHADE = "F2F2F2"
HEADER_SHADE = "1F3864"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

META_FIELDS = ["일시", "장소", "작성자", "참석자", "불참자", "안건"]
ACTION_COLS = ["할 일", "담당자", "기한", "상태"]
ACTION_WIDTHS = [4600, 2200, 1600, 960]  # twips, matches reference proportions

TABLE_WIDTH = 9360  # twips


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tcPr.append(shd)


def set_table_borders(table, color=BORDER_GRAY, size=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): str(size), qn("w:color"): color,
        })
        borders.append(el)
    tblPr.append(borders)


def set_col_widths(table, widths_twips):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_twips):
            cell.width = Twips(w)


def add_run(paragraph, text, size=10, bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def build(data, out_path):
    doc = Document()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    add_run(p, data["title"], size=16, bold=True, color=NAVY)

    # Subtitle with bottom border
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    if data.get("subtitle"):
        add_run(p, data["subtitle"], size=10, color=GRAY)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "12", qn("w:color"): "1F3864", qn("w:space"): "8",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Meta info table
    meta = data.get("meta", {})
    table = doc.add_table(rows=len(META_FIELDS), cols=2)
    set_table_borders(table)
    set_col_widths(table, [1800, 7560])
    for i, field in enumerate(META_FIELDS):
        label_cell, value_cell = table.rows[i].cells
        set_cell_shading(label_cell, LABEL_SHADE)
        lp = label_cell.paragraphs[0]
        add_run(lp, field, size=10, bold=True)
        vp = value_cell.paragraphs[0]
        add_run(vp, meta.get(field, ""), size=10)

    doc.add_paragraph()

    # Agenda items
    for item in data.get("agenda_items", []):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(5)
        add_run(h, item["heading"], size=12, bold=True, color=NAVY)

        if item.get("presenter"):
            pr = doc.add_paragraph()
            pr.paragraph_format.space_after = Pt(3)
            add_run(pr, f"발표: {item['presenter']}", size=9, italic=True, color=GRAY)

        for bullet in item.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet")
            add_run(bp, bullet, size=10)

    # Decision / action items
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(15)
    h.paragraph_format.space_after = Pt(6)
    add_run(h, "결정 사항 및 액션 아이템", size=12, bold=True, color=NAVY)

    action_items = data.get("action_items", [])
    table = doc.add_table(rows=1 + len(action_items), cols=4)
    set_table_borders(table)
    set_col_widths(table, ACTION_WIDTHS)
    for i, col in enumerate(ACTION_COLS):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_SHADE)
        cp = cell.paragraphs[0]
        add_run(cp, col, size=9, bold=True, color=WHITE)
    for r, ai in enumerate(action_items, start=1):
        values = [ai.get("task", ""), ai.get("owner", ""), ai.get("due", ""), ai.get("status", "")]
        for c, val in enumerate(values):
            cp = table.rows[r].cells[c].paragraphs[0]
            add_run(cp, val, size=9)

    # Next meeting
    if data.get("next_meeting"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(15)
        add_run(p, f"다음 회의: {data['next_meeting']}", size=9, color=GRAY)

    # Self-verification checklist
    checklist = data.get("checklist", [])
    if checklist:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(15)
        h.paragraph_format.space_after = Pt(5)
        add_run(h, "자체 검증 체크리스트", size=12, bold=True, color=NAVY)
        for entry in checklist:
            cp = doc.add_paragraph(style="List Bullet")
            mark = "✅" if entry.get("passed") else "⚠️"
            add_run(cp, f"{mark} {entry.get('text', '')}", size=10)

    doc.save(out_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 build_docx.py <input.json> <output.docx>", file=sys.stderr)
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        data = json.load(f)
    build(data, sys.argv[2])
    print(f"Wrote {sys.argv[2]}")


if __name__ == "__main__":
    main()
